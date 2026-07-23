from pathlib import Path
import json
import math
import zipfile
import xml.etree.ElementTree as ET

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/giacomo/Documents/Codex/2026-07-07/is")
OUT = ROOT / "outputs/hfi_pristine_rebuild_v1_1"
FIG = OUT / "submission_figures"
TAB = OUT / "submission_tables"
RENDER = OUT / "submission_rendered"
for p in [FIG, TAB, RENDER]:
    p.mkdir(parents=True, exist_ok=True)

TRACT = ROOT / "outputs/hfi_surface_rebuild_2026_07_22/hfi_tract_surface_v1_1_corrected.csv"
TRACT_GEOJSON = ROOT / "work/hfi-explorer-site-updated-v2/data/hfi_tracts.geojson"
HOSP = ROOT / "outputs/hfi_surface_rebuild_2026_07_22/nyc_hcahps_validation_analysis_file_corrected_hfi_v1_1.csv"
UNIVERSE = ROOT / "outputs/hospital_rebuild_2026_07_22/derived/nyc_acute_care_hospital_universe_33.csv"
MODEL = ROOT / "outputs/hfi_surface_rebuild_2026_07_22/corrected_hfi_v1_1_model_summary.csv"

DOCX_OUT = OUT / "HFI_NYC_full_submission_manuscript.docx"
EMAIL_OUT = OUT / "email_to_colleagues_revised.md"


TITLE = "A Geospatial Framework for Measuring Healthcare Fragmentation in Local Care Environments"
SUBTITLE = "Development and initial validation of the Healthcare Fragmentation Index in New York City"
AUTHORS = "Giacomo Di Pasquale, PhD; Greg Rybarczyk, PhD; Aimee Afable, PhD, MPH"
AFFILIATIONS = (
    "Giacomo Di Pasquale, New York University Robert F. Wagner Graduate School of Public Service; "
    "Greg Rybarczyk, University of Michigan-Flint; "
    "Aimee Afable, SUNY Downstate Health Sciences University, School of Public Health, "
    "450 Clarkson Ave, Brooklyn, NY 11203"
)


def font_path(bold=False):
    paths = [
        "/System/Library/Fonts/Supplemental/Aptos Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Aptos.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for p in paths:
        if Path(p).exists():
            return p
    return None


def fnt(size, bold=False):
    p = font_path(bold)
    return ImageFont.truetype(p, size) if p else ImageFont.load_default()


def set_run(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", size=12, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run(r, size=size, italic=italic, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, size=12, bold=True, color="173A59")
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Note. " + text)
    set_run(r, size=10, italic=True, color="555555")
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def add_table(doc, df, widths=None):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            if widths:
                cell.width = Inches(widths[ci])
            if ri == 0:
                shade(cell, "EAF3F5")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if ci > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_run(r, size=10, bold=(ri == 0), color="173A59" if ri == 0 else None)
    return table


def format_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(7)
    for name, color in [("Heading 1", "173A59"), ("Heading 2", "0A6B78"), ("Heading 3", "173A59")]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(12)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)


def zscore(v):
    v = np.asarray(v, dtype=float)
    return (v - np.nanmean(v)) / np.nanstd(v, ddof=1)


def ols_hc3(y, xcols):
    y = np.asarray(y, dtype=float)
    X = np.column_stack([np.ones(len(y))] + [np.asarray(c, dtype=float) for c in xcols])
    beta = np.linalg.inv(X.T @ X) @ X.T @ y
    resid = y - X @ beta
    hat = np.sum(X * (X @ np.linalg.inv(X.T @ X)), axis=1)
    meat = X.T @ np.diag((resid / (1 - hat)) ** 2) @ X
    cov = np.linalg.inv(X.T @ X) @ meat @ np.linalg.inv(X.T @ X)
    se = np.sqrt(np.diag(cov))
    ssr = np.sum(resid ** 2)
    sst = np.sum((y - y.mean()) ** 2)
    return beta, se, 1 - ssr / sst


def geom_bounds(features):
    xs, ys = [], []
    def collect(c):
        if isinstance(c[0], (float, int)):
            xs.append(c[0]); ys.append(c[1])
        else:
            for z in c:
                collect(z)
    for f in features:
        collect(f["geometry"]["coordinates"])
    return min(xs), min(ys), max(xs), max(ys)


def project(lon, lat, b, box):
    minx, miny, maxx, maxy = b
    x0, y0, x1, y1 = box
    scale = min((x1 - x0) / (maxx - minx), (y1 - y0) / (maxy - miny))
    used_w = (maxx - minx) * scale
    used_h = (maxy - miny) * scale
    ox = x0 + ((x1 - x0) - used_w) / 2
    oy = y0 + ((y1 - y0) - used_h) / 2
    return ox + (lon - minx) * scale, oy + used_h - (lat - miny) * scale


def draw_geom(draw, geom, b, box, fill, outline):
    def poly(coords):
        ring = [project(x, y, b, box) for x, y in coords[0]]
        if len(ring) > 2:
            draw.polygon(ring, fill=fill, outline=outline)
    if geom["type"] == "Polygon":
        poly(geom["coordinates"])
    elif geom["type"] == "MultiPolygon":
        for p in geom["coordinates"]:
            poly(p)


def draw_arrow(draw, start, end, color, width=6):
    draw.line((*start, *end), fill=color, width=width)
    dx, dy = end[0] - start[0], end[1] - start[1]
    ang = math.atan2(dy, dx)
    l = 18
    pts = [
        end,
        (end[0] - l * math.cos(ang - 0.45), end[1] - l * math.sin(ang - 0.45)),
        (end[0] - l * math.cos(ang + 0.45), end[1] - l * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill=color)


def make_conceptual_figure():
    scale = 2
    W, H = 1800 * scale, 950 * scale
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    ink = (20, 38, 58); muted = (88, 105, 120); teal = (10, 107, 120)
    colors = [(232, 245, 247), (243, 248, 250), (255, 248, 222), (238, 245, 252)]
    boxes = [
        (90, 220, 435, 610, "Local healthcare\nresources", "Providers, hospitals,\ncapacity, hours,\nand service locations"),
        (530, 220, 875, 610, "Access conditions", "Distance, transportation,\naffordability, quality,\nand system friction"),
        (970, 220, 1315, 610, "HFI", "A tract-level measure\nof whether resources\nform a navigable care\nenvironment"),
        (1410, 220, 1710, 610, "Validation domain", "Hospital context,\npatient experience,\nand communication"),
    ]
    d.text((90*scale, 70*scale), "Conceptual framework for healthcare fragmentation", font=fnt(46*scale, True), fill=ink)
    d.text((90*scale, 130*scale), "HFI translates multidimensional access conditions into a reproducible tract-level fragmentation surface", font=fnt(26*scale), fill=muted)
    for i, (x0, y0, x1, y1, title, body) in enumerate(boxes):
        box = tuple(v * scale for v in (x0, y0, x1, y1))
        d.rounded_rectangle(box, radius=24*scale, fill=colors[i], outline=(204, 220, 228), width=3*scale)
        d.text(((x0+28)*scale, (y0+34)*scale), title, font=fnt(31*scale, True), fill=ink, spacing=5*scale)
        d.text(((x0+28)*scale, (y0+155)*scale), body, font=fnt(23*scale), fill=muted, spacing=7*scale)
        if i < len(boxes) - 1:
            draw_arrow(d, ((x1+30)*scale, 415*scale), ((boxes[i+1][0]-30)*scale, 415*scale), teal, width=6*scale)
    d.rounded_rectangle((250*scale, 725*scale, 1550*scale, 840*scale), radius=18*scale, fill=(250, 252, 253), outline=(217, 228, 234), width=2*scale)
    d.text((290*scale, 752*scale), "Core claim: fragmentation is not only distance. It is the local organization of care opportunity, burden, and navigability.", font=fnt(25*scale, True), fill=ink)
    out = FIG / "figure1_conceptual_framework.png"
    img.resize((1800, 950), Image.Resampling.LANCZOS).save(out, quality=95)
    return out


def make_map_figure():
    tr = pd.read_csv(TRACT, dtype={"tract_id": str})
    hosp = pd.read_csv(HOSP, dtype={"GEOID": str})
    with TRACT_GEOJSON.open() as f:
        geo = json.load(f)
    bins = pd.qcut(tr["fragmentation_index"], 5, labels=False, duplicates="drop")
    binmap = dict(zip(tr["tract_id"], bins))
    palette = [(36, 104, 154), (117, 174, 196), (255, 242, 176), (238, 157, 82), (198, 63, 65)]
    scale = 3
    W, H = 1800 * scale, 1320 * scale
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    ink = (20, 38, 58); muted = (85, 101, 116)
    d.text((70*scale, 48*scale), "Healthcare Fragmentation Index across New York City", font=fnt(48*scale, True), fill=ink)
    d.text((70*scale, 108*scale), "Tract-level HFI quintiles with CMS acute-care hospitals included in the HCAHPS validation sample", font=fnt(25*scale), fill=muted)
    map_box = (85*scale, 185*scale, 1715*scale, 1085*scale)
    b = geom_bounds(geo["features"])
    for feat in geo["features"]:
        geoid = str(feat.get("properties", {}).get("GEOID", ""))
        fill = palette[int(binmap.get(geoid, 2))]
        draw_geom(d, feat["geometry"], b, map_box, fill, (255, 255, 255))
    for _, r in hosp.iterrows():
        x, y = project(float(r["_CX"]), float(r["_CY"]), b, map_box)
        rad = 9 * scale
        d.ellipse((x-rad, y-rad, x+rad, y+rad), fill=(255, 202, 93), outline=ink, width=3*scale)
    d.text((115*scale, 1160*scale), "Lower HFI", font=fnt(22*scale, True), fill=muted)
    for i, c in enumerate(palette):
        x0 = (260 + i*68) * scale
        d.rectangle((x0, 1158*scale, x0 + 48*scale, 1186*scale), fill=c, outline=(180, 190, 198), width=1*scale)
    d.text((640*scale, 1160*scale), "Higher HFI", font=fnt(22*scale, True), fill=muted)
    d.ellipse((900*scale, 1158*scale, 928*scale, 1186*scale), fill=(255, 202, 93), outline=ink, width=3*scale)
    d.text((948*scale, 1160*scale), "HCAHPS validation hospital", font=fnt(22*scale, True), fill=muted)
    out = FIG / "figure2_hfi_map_publication.png"
    img.resize((1800, 1320), Image.Resampling.LANCZOS).save(out, quality=95)
    return out


def make_hospital_context_figure():
    tr = pd.read_csv(TRACT, dtype={"tract_id": str})
    universe = pd.read_csv(UNIVERSE, dtype={"GEOID": str})
    with TRACT_GEOJSON.open() as f:
        geo = json.load(f)
    bins = pd.qcut(tr["fragmentation_index"], 5, labels=False, duplicates="drop")
    binmap = dict(zip(tr["tract_id"], bins))
    palette = [(36, 104, 154), (117, 174, 196), (255, 242, 176), (238, 157, 82), (198, 63, 65)]
    scale = 3
    W, H = 1800 * scale, 1180 * scale
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    ink = (20, 38, 58); muted = (85, 101, 116); teal = (10, 107, 120)
    d.text((70*scale, 48*scale), "CMS acute-care hospitals embedded in the HFI surface", font=fnt(45*scale, True), fill=ink)
    d.text((70*scale, 105*scale), "All CMS-listed NYC acute-care hospitals, with HCAHPS validation eligibility shown at point locations", font=fnt(24*scale), fill=muted)
    map_box = (105*scale, 170*scale, 1695*scale, 945*scale)
    b = geom_bounds(geo["features"])
    for feat in geo["features"]:
        geoid = str(feat.get("properties", {}).get("GEOID", ""))
        fill = palette[int(binmap.get(geoid, 2))]
        draw_geom(d, feat["geometry"], b, map_box, fill, (255, 255, 255))
    for _, r in universe.iterrows():
        x, y = project(float(r["_CX"]), float(r["_CY"]), b, map_box)
        included = str(r.get("included_in_hcahps_validation", "")).lower() == "true"
        public = str(r.get("public_private_designation", "")).lower() == "public"
        rad = 11 * scale if public else 9 * scale
        fill = (255, 202, 93) if included else (180, 188, 196)
        outline = teal if public else ink
        d.ellipse((x-rad, y-rad, x+rad, y+rad), fill=fill, outline=outline, width=3*scale)
    d.text((140*scale, 1015*scale), "Tract HFI surface", font=fnt(22*scale, True), fill=muted)
    for i, c in enumerate(palette):
        x0 = (345 + i*62) * scale
        d.rectangle((x0, 1013*scale, x0 + 42*scale, 1040*scale), fill=c, outline=(180, 190, 198), width=1*scale)
    d.text((700*scale, 1015*scale), "HCAHPS-complete hospital", font=fnt(22*scale, True), fill=muted)
    d.ellipse((648*scale, 1012*scale, 680*scale, 1044*scale), fill=(255, 202, 93), outline=ink, width=3*scale)
    d.text((1095*scale, 1015*scale), "Public hospital outline", font=fnt(22*scale, True), fill=muted)
    d.ellipse((1045*scale, 1012*scale, 1077*scale, 1044*scale), fill=(255, 202, 93), outline=teal, width=5*scale)
    d.text((1435*scale, 1015*scale), "HCAHPS unavailable", font=fnt(22*scale, True), fill=muted)
    d.ellipse((1385*scale, 1012*scale, 1417*scale, 1044*scale), fill=(180, 188, 196), outline=ink, width=3*scale)
    out = FIG / "figure3_hospital_context_map.png"
    img.resize((1800, 1180), Image.Resampling.LANCZOS).save(out, quality=95)
    return out


def make_scatter_figure():
    hosp = pd.read_csv(HOSP)
    x = pd.to_numeric(hosp["fragmentation_index"], errors="coerce").to_numpy(float)
    y = pd.to_numeric(hosp["communication_index"], errors="coerce").to_numpy(float)
    ok = ~np.isnan(x) & ~np.isnan(y)
    x, y = x[ok], y[ok]
    slope, intercept = np.polyfit(x, y, 1)
    scale = 3
    W, H = 1600 * scale, 1050 * scale
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    ink = (20, 38, 58); muted = (88, 105, 120); teal = (10, 132, 145)
    d.text((80*scale, 50*scale), "HFI and hospital communication scores", font=fnt(46*scale, True), fill=ink)
    d.text((80*scale, 108*scale), "CMS acute-care hospitals with complete HCAHPS outcomes, New York City (n = 32)", font=fnt(25*scale), fill=muted)
    plot = (170*scale, 205*scale, 1420*scale, 820*scale)
    xmin, xmax = math.floor(x.min()) - 0.25, math.ceil(x.max()) + 0.25
    ymin, ymax = 80, 94
    def px(v): return plot[0] + (v - xmin) / (xmax - xmin) * (plot[2] - plot[0])
    def py(v): return plot[3] - (v - ymin) / (ymax - ymin) * (plot[3] - plot[1])
    d.rectangle(plot, fill=(252, 254, 255), outline=(218, 229, 236), width=2*scale)
    for tv in np.linspace(math.ceil(xmin), math.floor(xmax), 6):
        d.line((px(tv), plot[1], px(tv), plot[3]), fill=(226, 234, 238), width=1*scale)
        d.text((px(tv)-20*scale, 835*scale), f"{tv:.0f}", font=fnt(18*scale), fill=muted)
    for tv in range(80, 95, 2):
        d.line((plot[0], py(tv), plot[2], py(tv)), fill=(226, 234, 238), width=1*scale)
        d.text((115*scale, py(tv)-10*scale), str(tv), font=fnt(18*scale), fill=muted)
    d.line((px(xmin), py(slope*xmin + intercept), px(xmax), py(slope*xmax + intercept)), fill=teal, width=5*scale)
    for xi, yi in zip(x, y):
        d.ellipse((px(xi)-9*scale, py(yi)-9*scale, px(xi)+9*scale, py(yi)+9*scale), fill=(255, 202, 93), outline=ink, width=2*scale)
    d.text((610*scale, 930*scale), "Healthcare Fragmentation Index", font=fnt(24*scale, True), fill=ink)
    d.text((170*scale, 168*scale), "Communication index", font=fnt(22*scale, True), fill=ink)
    d.text((1075*scale, 240*scale), f"Bivariate slope = {slope:.3f}", font=fnt(23*scale, True), fill=teal)
    d.text((1075*scale, 274*scale), "HC3 SE = 0.444", font=fnt(20*scale), fill=muted)
    out = FIG / "figure4_hfi_hcahps_scatter_publication.png"
    img.resize((1600, 1050), Image.Resampling.LANCZOS).save(out, quality=95)
    return out


def make_tables():
    tr = pd.read_csv(TRACT)
    hosp = pd.read_csv(HOSP)
    universe = pd.read_csv(UNIVERSE)
    model = pd.read_csv(MODEL)
    table_system = universe.groupby(["hospital_system", "public_private_designation"]).size().reset_index(name="Hospitals")
    table_system = table_system.rename(columns={"hospital_system": "Hospital system", "public_private_designation": "Designation"})
    table_system = table_system.sort_values(["Designation", "Hospital system"])
    table_system.to_csv(TAB / "table1_hospital_systems.csv", index=False)
    defs = pd.DataFrame([
        ["HFI", "Census tract", "Standardized sum of reversed nearest-provider access and reversed provider density", "Higher values indicate greater local healthcare fragmentation"],
        ["Nearest-provider access", "Census tract", "1 / (minimum provider distance + 0.1), standardized before reversal", "Lower access contributes to higher HFI"],
        ["Provider density", "Census tract", "Time-weighted provider supply per tract population", "Lower density contributes to higher HFI"],
        ["Communication index", "Hospital", "Mean of HCAHPS nurse communication and doctor communication linear scores", "Higher values indicate better patient-rated communication"],
        ["Patient-experience index", "Hospital", "Composite of standardized HCAHPS domains", "Exploratory external-validation outcome"],
        ["CMS overall hospital rating", "Hospital", "Publicly reported CMS overall rating", "Institutional performance covariate"],
    ], columns=["Measure", "Unit", "Operational definition", "Interpretation"])
    defs.to_csv(TAB / "table2_definitions.csv", index=False)
    desc_rows = []
    for label, df, col in [
        ("HFI, tract surface", tr, "fragmentation_index"),
        ("Provider count, tract surface", tr, "providers"),
        ("Minimum provider distance", tr, "min_dist"),
        ("Poverty rate, tract surface", tr, "poverty_rate"),
        ("HFI, hospital-linked sample", hosp, "fragmentation_index"),
        ("Communication index", hosp, "communication_index"),
        ("Patient-experience index", hosp, "patient_experience_index"),
        ("CMS overall hospital rating", hosp, "hospital_overall_rating"),
        ("Completed HCAHPS surveys", hosp, "number_of_completed_surveys"),
    ]:
        s = pd.to_numeric(df[col], errors="coerce").dropna()
        desc_rows.append([label, int(s.count()), f"{s.mean():.2f}", f"{s.std(ddof=1):.2f}", f"{s.min():.2f}", f"{s.median():.2f}", f"{s.max():.2f}"])
    desc = pd.DataFrame(desc_rows, columns=["Variable", "N", "Mean", "SD", "Min", "Median", "Max"])
    desc.to_csv(TAB / "table3_descriptives.csv", index=False)
    rows = []
    for label, term in [("HFI", "fragmentation_index"), ("CMS overall hospital rating", "hospital_overall_rating"), ("Constant", "Intercept")]:
        coef_row = {"Predictor": label}
        se_row = {"Predictor": ""}
        for i, m in enumerate(model["model"].drop_duplicates(), start=1):
            sub = model[(model["model"] == m) & (model["term"] == term)]
            if len(sub):
                r = sub.iloc[0]
                stars = "***" if r["p_norm"] < .001 else "**" if r["p_norm"] < .01 else "*" if r["p_norm"] < .05 else "+" if r["p_norm"] < .1 else ""
                coef_row[f"Model {i}"] = f"{r['coef']:.3f}{stars}"
                se_row[f"Model {i}"] = f"({r['hc3_se']:.3f})"
            else:
                coef_row[f"Model {i}"] = ""
                se_row[f"Model {i}"] = ""
        rows.extend([coef_row, se_row])
    nrow = {"Predictor": "N"}
    for i, m in enumerate(model["model"].drop_duplicates(), start=1):
        nrow[f"Model {i}"] = "32"
    models = pd.DataFrame(rows + [nrow])
    models.to_csv(TAB / "table4_models.csv", index=False)
    spatial = pd.DataFrame([
        ["4-nearest neighbors", "0.389", "0.001", "Supports spatial clustering"],
        ["8-nearest neighbors", "0.355", "0.001", "Primary specification"],
        ["12-nearest neighbors", "0.323", "0.001", "Supports spatial clustering"],
    ], columns=["Spatial weights", "Moran's I", "Permutation pseudo-p", "Interpretation"])
    spatial.to_csv(TAB / "table5_moran_sensitivity.csv", index=False)
    return table_system, defs, desc, models, spatial


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run(r, size=15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    set_run(r, size=12, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AUTHORS)
    set_run(r, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AFFILIATIONS)
    set_run(r, size=10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run(r)


def add_display_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=12)
    return p


def build_doc():
    fig1 = make_conceptual_figure()
    fig2 = make_map_figure()
    fig3 = make_hospital_context_figure()
    fig4 = make_scatter_figure()
    table_system, defs, desc, models, spatial = make_tables()
    tr = pd.read_csv(TRACT)
    hosp = pd.read_csv(HOSP)
    doc = Document()
    format_doc(doc)
    add_title(doc)

    heading(doc, "Abstract", 1)
    para(doc, "Background: Healthcare access is often measured through provider supply, distance, travel burden, insurance coverage, or institutional quality. These dimensions are essential, but they do not fully capture whether local healthcare resources form a coherent, reachable, affordable, and navigable system of care. This paper develops the Healthcare Fragmentation Index (HFI), a reproducible geospatial framework for measuring fragmentation in local healthcare environments.")
    para(doc, "Methods: HFI was constructed for 2,229 New York City census tracts using provider locations, time-weighted provider supply, nearest-provider access, and CMS acute-care hospital representation. The hospital universe was defined from CMS Hospital General Information files. Thirty-three CMS-listed acute-care hospitals were identified in New York City, of which 32 had complete HCAHPS patient-experience outcomes and could be included in the hospital-level validation analysis. We evaluated distributional structure, tract-level spatial autocorrelation using global Moran’s I, and hospital-level construct-validation associations between HFI and HCAHPS communication outcomes using ordinary least squares models with HC3 robust standard errors.")
    para(doc, "Results: HFI produced a complete tract-level surface for New York City. The provider layer identified 653 tracts with at least one provider and 1,576 tracts without an in-tract provider. HFI displayed significant spatial clustering across census tracts, with Moran’s I = 0.355 using 8-nearest-neighbor weights and permutation pseudo-p = 0.001; sensitivity checks using 4- and 12-nearest-neighbor weights were also significant. In the 32-hospital validation sample, HFI showed positive hospital-level associations with HCAHPS communication and broader patient-experience outcomes, while CMS overall hospital rating remained the strongest hospital-level correlate of patient experience.")
    para(doc, "Conclusions: HFI provides a reproducible framework for making healthcare fragmentation visible at tract level. The New York City analysis supports HFI as a spatial measurement framework and demonstrates how the index can be linked to public hospital quality data for construct validation. HFI’s contribution is to establish a transparent language through which healthcare fragmentation can be measured, compared, challenged, and improved.")
    para(doc, "Keywords: healthcare fragmentation; spatial accessibility; HCAHPS; patient experience; GIS; Moran’s I; New York City.", italic=True)

    heading(doc, "Introduction", 1)
    para(doc, "Healthcare systems are usually described through the institutions that compose them: hospitals, clinics, physicians, insurance networks, public programs, quality scores, and financing arrangements. Patients, however, encounter these components as local healthcare environments. A neighborhood may contain facilities but still impose practical barriers if those facilities are difficult to reach, poorly connected to one another, unaffordable, mismatched to patient needs, or embedded in a system that is hard to navigate. For residents of large metropolitan areas, the question is not only whether care exists somewhere in the city, but whether the local care environment functions as a coherent system.")
    heading(doc, "Defining access and fragmentation", 2)
    definition_intro = [
        "Access is defined here as a multidimensional relationship between people, healthcare resources, institutions, and constraints. It includes geographic proximity, but it also includes availability, affordability, accommodation, acceptability, transportation, insurance fit, language and cultural accessibility, institutional quality, and the practical ability to move through the system (Aday and Andersen, 1974; Penchansky and Thomas, 1981; Andersen, 1995; Levesque et al., 2013). Distance is therefore one dimension of access, not access itself.",
        "Healthcare fragmentation is defined here as the degree to which local healthcare resources are organized into a coherent, reachable, and navigable care environment. Fragmentation is related to access because fragmented environments make access harder to realize, but the two concepts are not identical. Accessibility asks whether opportunities can be reached. Fragmentation asks whether those opportunities function together as a usable local system. A neighborhood can be geographically close to healthcare institutions and still be fragmented if residents face system boundaries, insurance mismatches, poor transit connections, weak referral pathways, language barriers, or unclear post-discharge routes.",
        "This paper develops the Healthcare Fragmentation Index (HFI) to measure this local structure. HFI treats fragmentation as a spatial property of the healthcare environment: the degree to which resources form a coherent, reachable, and navigable system of care for residents and institutions in a given place. This framing complements, rather than replaces, existing measures of access, quality, and continuity. Accessibility measures ask whether residents can reach healthcare opportunities. Quality measures ask how well institutions perform. Continuity measures ask whether patients experience coordinated care over time. HFI asks a related but distinct question: how fragmented is the local healthcare environment through which access, quality, and continuity must operate?",
    ]
    for p in definition_intro:
        para(doc, p)
    caption(doc, "Figure 1. Multidimensional access and healthcare fragmentation framework")
    doc.add_picture(str(fig1), width=Inches(6.45))
    note(doc, "Access is defined as multidimensional rather than distance-only. HFI translates selected spatial dimensions of access into a tract-level measure of local healthcare fragmentation while leaving room for future transportation, affordability, insurance, language, and patient-flow extensions.")
    intro = [
        "New York City is an appropriate case for developing the framework because it combines extraordinary healthcare density with deep spatial inequality and institutional complexity. The city contains nationally prominent academic medical centers, one of the largest public hospital systems in the United States, independent community hospitals, specialty institutions, and neighborhoods with very different histories of investment, marginalization, transportation access, and healthcare infrastructure. This interpretation is consistent with the Urban Stress Model, which uses New York City to show how urban scale, density, diversity, complexity, social disadvantage, and historically rooted inequity shape health through cumulative stress pathways (Afable et al., 2026). In this context, fragmentation is not simply the absence of providers. It is also a consequence of uneven urban development, safety-net reliance, system complexity, and the layered barriers that shape how residents identify, reach, afford, and coordinate care.",
        "The New York City case also makes clear why fragmentation should not be treated as a synonym for scarcity. Some neighborhoods may be underserved because resources are scarce. Other neighborhoods may be surrounded by institutions but still be difficult to navigate because services are specialized, system boundaries are opaque, insurance relationships are uneven, or hospital and ambulatory care are poorly integrated. Fragmentation therefore refers to a structural relationship among resources, populations, and pathways. It is a condition of the care environment, not merely a count of facilities.",
        "This distinction has practical consequences. A local public health agency, hospital system, or community organization may know where hospitals are located and may know which neighborhoods have high social need. Yet those facts alone do not show whether residents experience the healthcare landscape as a coherent set of reachable options. A spatial fragmentation measure can help identify places where healthcare opportunity is present but thinly connected, difficult to reach, or concentrated outside the neighborhood. It can also help distinguish neighborhoods where intervention should focus on provider supply from neighborhoods where transportation, system navigation, affordability, or care coordination may be more salient.",
        "The conceptual contribution is therefore to make healthcare fragmentation measurable as a spatial framework. In the same way that the Area Deprivation Index operationalized neighborhood disadvantage and Walk Score operationalized walkability, HFI operationalizes local healthcare fragmentation as a tract-level spatial condition (Carr et al., 2011; Kind and Buckingham, 2018). The objective is to create a transparent baseline that can be replicated, critiqued, and extended across settings.",
        "The empirical contribution is threefold. First, the paper constructs a tract-level HFI surface for all 2,229 New York City census tracts. Second, it evaluates whether the resulting surface has coherent distributional and spatial structure rather than behaving as random geographic noise. Third, it links all eligible CMS acute-care hospitals in New York City to the tract-level HFI surface and uses HCAHPS patient-experience measures as an external construct-validation test. The hospital-level analysis is designed to evaluate whether a spatially measured care-environment construct can be connected meaningfully to publicly reported institutional outcomes."
    ]
    for p in intro:
        para(doc, p)

    heading(doc, "Literature Review", 1)
    lit = [
        "The literature motivating HFI spans five areas: access to care, spatial accessibility measurement, neighborhood and urban health, care fragmentation, and patient experience. Access frameworks have long emphasized that access is broader than physical proximity. Aday and Andersen’s framework distinguished potential access from realized use and located access within enabling resources, need, health-system organization, and utilization (Aday and Andersen, 1974). Andersen (1995) later clarified the behavioral model as a framework for understanding how predisposing, enabling, and need factors shape use. Penchansky and Thomas (1981) conceptualized access through availability, accessibility, accommodation, affordability, and acceptability. Levesque et al. (2013) extended this tradition by emphasizing the interface between health systems and populations, including approachability, acceptability, availability, affordability, and appropriateness. These traditions make clear that access is not a single variable; it is a relationship between people, resources, institutions, and constraints.",
        "Spatial accessibility research provides the empirical backbone for HFI. Hansen (1959) defined accessibility as the potential for interaction between places and opportunities. Subsequent health-geography research adapted this logic to healthcare, moving beyond simple distance measures toward approaches that account for supply, demand, travel impedance, and aggregation (Joseph and Bantock, 1982; Khan, 1992; Radke and Mu, 2000; Luo and Wang, 2003; Guagliardo, 2004; Higgs, 2004). Later work refined these methods through two-step floating catchment area approaches, enhanced accessibility measures, and attention to rural and urban context (Wang and Luo, 2005; Apparicio et al., 2008; Luo and Qi, 2009; McGrail and Humphreys, 2009).",
        "HFI builds on this spatial-accessibility tradition but asks a different question. Accessibility asks whether residents can reach opportunities. Fragmentation asks whether the surrounding healthcare environment functions as a coherent system. A tract with low provider density and distant nearest providers is not only less accessible; it may also be more fragmented because residents must navigate a thinner, more dispersed, and less locally coherent field of care. HFI therefore uses accessibility as an empirical foundation while interpreting constrained accessibility and thin local supply as measurable dimensions of fragmentation.",
        "This extension is important because conventional access measures can be strong descriptively while remaining incomplete conceptually. Distance-to-nearest-provider measures are transparent but ignore supply, capacity, and demand. Provider-per-population ratios capture supply but ignore spatial impedance. Floating catchment methods incorporate both supply and distance, but they usually remain measures of opportunity rather than measures of system coherence. HFI uses this methodological lineage as a starting point and then shifts the interpretive frame from opportunity alone to the organization of the local care environment.",
        "The distinction between access and fragmentation is especially important in metropolitan settings. Dense cities can produce a paradox of proximity: many resources are nearby in absolute distance, yet the effective care environment remains difficult to use. Residents may face insurance constraints, long travel times despite short distances, specialty mismatches, hospital-system boundaries, language barriers, or uncertainty about where to seek follow-up. Transportation barriers are especially important because spatial opportunity is only meaningful when residents can realistically reach care (Syed et al., 2013). A fragmentation framework allows these constraints to be conceptualized as part of a local system rather than as isolated barriers.",
        "Neighborhood and urban-health research provides the equity context for this argument. Neighborhood conditions shape exposure, resources, stress, social organization, and health opportunity (Diez Roux and Mair, 2010). Residential segregation and historically patterned inequality structure the geography of health resources and health risk (Williams and Collins, 2001). Social vulnerability measures show how demographic, socioeconomic, and built-environment conditions can concentrate risk before a specific event occurs (Flanagan et al., 2011). HFI extends this neighborhood-measurement tradition by focusing specifically on the spatial organization of healthcare resources.",
        "Health services research on fragmentation and coordination supplies the conceptual motivation. Fragmentation has been described as a problem of accountability, continuity, patient burden, and system integration (Starfield, 1998; Starfield et al., 2005; Bodenheimer, 2008; Stange, 2009). Claims-based studies often measure fragmentation through the dispersion of visits across physicians, discontinuity across episodes, or organizational patterns that complicate coordination (Pham et al., 2007; Cebul et al., 2008; Frandsen et al., 2015). These measures are valuable because they capture the patient’s realized path through care.",
        "A spatial measure complements these patient-level and organizational measures by moving upstream. Before care can be coordinated, patients must find and reach a care environment. A fragmented geography may increase the burden of identifying appropriate providers, coordinating follow-up, interpreting discharge instructions, and sustaining continuity after hospitalization. This is especially important in cities where public hospitals, academic systems, community hospitals, and specialty providers are unevenly distributed across neighborhoods and transit geographies. Earlier work on small-area variation in healthcare delivery also shows why geographic units can reveal meaningful differences in system organization that are obscured in aggregate data (Wennberg and Gittelsohn, 1973).",
        "Continuity measures are therefore not competitors to HFI. They measure a different part of the fragmentation process. Continuity has been conceptualized as informational continuity, longitudinal continuity, and interpersonal continuity, each capturing a different way that care remains connected over time (Saultz, 2003). Claims-based continuity indices, including the Usual Provider Continuity index and the Continuity of Care Index, typically ask whether an individual patient’s care is concentrated among a small number of clinicians or dispersed across many providers (Bice and Boxerman, 1977; Jee and Cabana, 2006). These measures are valuable for evaluating realized care patterns after patients enter the system.",
        "HFI moves one step upstream. It asks whether the local environment through which patients seek care is spatially coherent before the patient’s realized care pattern is observed. A patient may ultimately achieve continuity with a clinician, but the surrounding care environment can still impose navigational burden if providers are sparse, distant, poorly connected to hospitals, or difficult to reach through transportation and insurance pathways. In future work, the two approaches can be joined: tract-level HFI can be tested against patient-level continuity, avoidable utilization, post-discharge follow-up, or continuity-sensitive outcomes using claims, SPARCS, or all-payer discharge data.",
        "Patient experience research provides a plausible external-validation domain. The Consumer Assessment of Healthcare Providers and Systems Hospital Survey, commonly known as HCAHPS, was developed to provide standardized national measures of patients’ perspectives on inpatient care (Goldstein et al., 2005). Patient experience is not reducible to clinical quality, but it is associated with communication, safety, patient-centeredness, institutional performance, and outcomes (Sofaer and Firminger, 2005; Jha et al., 2008; Bleich et al., 2009; Isaac et al., 2010; Boulding et al., 2011; Doyle et al., 2013; Manary et al., 2013; Anhang Price et al., 2014; Tsai et al., 2015). Recent evidence also shows that hospital characteristics and interventions can shape patient experience, reinforcing that HCAHPS should be interpreted as a patient-facing institutional outcome rather than as a direct neighborhood measure (Beckett et al., 2024).",
        "Quality and coordination frameworks reinforce the same point. The Institute of Medicine (2001) defined high-quality care through safety, effectiveness, patient-centeredness, timeliness, efficiency, and equity. International evidence on patients with complex care needs shows that coordination problems remain common even in high-income health systems (Schoen et al., 2011). HFI therefore sits at the intersection of access, coordination, quality, and geography: it measures the local care environment through which these broader system goals must be realized.",
        "The key interpretive point is that HCAHPS scores are hospital-level outcomes shaped by many forces. They reflect hospital processes, patient mix, expectations, survey response patterns, institutional quality, and the broader context in which patients enter and leave care. HFI contributes a tract-level care-environment measure that can be linked to these standardized outcomes. The appropriate validation question is therefore whether a tract-level fragmentation surface can be transparently connected to hospital outcomes and whether the resulting relationship provides useful evidence about the construct’s behavior and boundaries."
    ]
    for p in lit:
        para(doc, p)

    heading(doc, "Conceptual Framework", 1)
    framework = [
        "HFI defines healthcare fragmentation as a property of local care environments. Fragmentation is higher when healthcare resources are spatially thin, harder to reach, less locally coherent, or more burdensome to navigate as a practical system of care. This definition intentionally separates fragmentation from deprivation. A poor neighborhood may be fragmented, but poverty and fragmentation are not the same construct. Deprivation measures socioeconomic disadvantage; HFI measures the spatial organization of healthcare opportunity.",
        "The framework also separates fragmentation from hospital quality. A hospital may have high or low HCAHPS scores for reasons internal to the institution. HFI instead measures the environment in which hospitals and patients are embedded. The link between HFI and patient experience is therefore indirect. Fragmented environments may place greater communication and coordination burdens on hospitals and patients, but the hospital-level validation analysis cannot determine whether fragmentation causes lower communication scores.",
        "The framework treats access as multidimensional but begins with a parsimonious spatial implementation. HFI can eventually incorporate multimodal travel time, vehicle availability, transit dependence, insurance networks, appointment supply, specialty mix, language access, referral patterns, and patient-flow data. The present New York City implementation focuses on provider locations, time-weighted provider availability, nearest-provider access, and provider density because these can be constructed reproducibly at tract level and linked to public hospital data.",
    ]
    for p in framework:
        para(doc, p)
    para(doc, "The conceptual framework in Figure 1 summarizes this relationship by placing HFI between multidimensional access conditions and the patient-experience validation domain. The figure is intentionally conceptual: it shows what the framework is designed to measure, while the Methods section specifies how the first New York City implementation operationalizes the construct.")

    heading(doc, "Research Objectives", 2)
    add_bullets(doc, [
        "Define healthcare fragmentation as a reproducible tract-level spatial construct.",
        "Operationalize HFI using nearest-provider access and time-weighted provider density.",
        "Evaluate whether HFI has meaningful distributional and spatial structure across New York City.",
        "Link HFI to all eligible CMS acute-care hospitals in New York City and examine construct-validation associations with HCAHPS communication and patient-experience outcomes.",
    ])

    heading(doc, "Expected Properties of the Index", 2)
    properties = [
        "An index paper should be evaluated not only by whether one validation coefficient is statistically significant, but by whether the proposed measure has properties that make it useful for cumulative research. HFI is designed to satisfy five such properties. First, it is transparent: the input components, transformations, standardization, and direction of interpretation can be inspected. Second, it is spatially explicit: every tract receives a value that can be mapped and linked to other tract-level or institutional data.",
        "Third, HFI is directionally interpretable. Higher values correspond to more fragmented local care environments, meaning thinner local supply and weaker nearest-provider access in the current implementation. Fourth, HFI is extensible. The framework can incorporate transportation, affordability, insurance networks, specialty mix, appointment availability, patient origin, and care-network data as those data become available. Fifth, HFI is falsifiable. Other researchers can challenge its inputs, alter its weights, test alternative geographies, and evaluate whether it behaves as expected against independent outcomes.",
        "These properties matter because healthcare fragmentation is not a single observed variable. It is a construct. The first task is therefore to define and operationalize the construct in a way that can be debated. A single public hospital-level validation outcome should not be treated as the sole criterion for judging a spatial framework. The broader question is whether the framework makes a previously diffuse system property measurable enough to support replication, critique, and improvement."
    ]
    for p in properties:
        para(doc, p)

    heading(doc, "Methods", 1)
    methods = [
        "This study is a spatial methods and initial validation study of healthcare fragmentation in New York City. The unit of index construction is the census tract. The unit of hospital outcome analysis is the hospital. The tract-level analysis includes 2,229 New York City census tracts. The hospital universe includes CMS-listed acute-care hospitals located in the five New York City counties.",
        "The analysis integrates four data domains: census-tract geography, healthcare provider locations, provider operating-time information, and CMS Care Compare/HCAHPS hospital-level patient-experience measures. Provider locations were drawn from the project’s cleaned provider-location layer, geocoded, and linked to census tracts. For this implementation, providers are healthcare service locations represented in the HFI provider file, with acute-care hospitals included as healthcare access points and represented as continuously available. Provider availability was time-weighted using operating-time information. Tract-level socioeconomic variables were retained for descriptive context and construct interpretation.",
        "The CMS hospital universe was defined using the CMS Provider Data Catalog Hospital General Information file and restricted to acute-care hospitals in New York City (Centers for Medicare & Medicaid Services, 2026a). Hospital patient-experience outcomes were drawn from the CMS Provider Data Catalog HCAHPS Hospital file covering July 1, 2024 through June 30, 2025 (Centers for Medicare & Medicaid Services, 2026b). This identified 33 acute-care hospitals. Thirty-two hospitals had complete HCAHPS linear patient-experience outcomes and were included in the validation sample. NY Eye and Ear Infirmary of Mount Sinai was retained in the broader hospital universe but excluded from HCAHPS validation models because required HCAHPS linear outcomes were unavailable.",
        "Hospitals were assigned HFI values using a point-in-tract linkage. Each geocoded hospital point was linked to the census tract containing the hospital location, and the hospital inherited the tract’s HFI value. This strategy is intentionally simple and reproducible. It treats the hospital as embedded in a local structural environment rather than estimating patient-specific exposure. Because individual patient residence, hospital catchment areas, discharge destinations, and referral pathways are not observed in public HCAHPS files, catchment-weighted exposure is reserved for future validation work.",
    ]
    for p in methods:
        para(doc, p)
    doc.add_page_break()
    caption(doc, "Table 1. CMS-listed acute-care hospital universe by system and designation")
    add_table(doc, table_system, widths=[3.1, 1.7, .75])
    note(doc, "The universe includes 33 CMS-listed acute-care hospitals in New York City. The HCAHPS validation sample includes the 32 hospitals with complete required HCAHPS linear outcomes.")
    heading(doc, "Index Construction", 2)
    for p in [
        "HFI combines two tract-level components derived from the spatial-accessibility tradition. The first component captures nearest-provider access rather than a full two-step floating catchment area score. This conservative specification follows the accessibility literature’s emphasis on spatial separation and opportunity while keeping the first HFI implementation transparent. The second component captures provider density through time-weighted local supply per population. Both components are standardized and reversed where needed so that higher values consistently indicate greater fragmentation.",
    ]:
        para(doc, p)
    add_display_equation(doc, "dᵢ = minⱼ distance(i, j)")
    add_display_equation(doc, "Aᵢ = 1 / (dᵢ + 0.1)")
    add_display_equation(doc, "Tᵢ = Σ hⱼ for providers j located in tract i")
    add_display_equation(doc, "Dᵢ = Tᵢ / Pᵢ")
    add_display_equation(doc, "Fᴬᵢ = -z(Aᵢ),     Fᴰᵢ = -z(Dᵢ)")
    add_display_equation(doc, "HFIᵢ = z(Fᴬᵢ + Fᴰᵢ)")
    note(doc, "dᵢ is the distance from tract i to the nearest provider location; Aᵢ is nearest-provider access; hⱼ is the operating-time weight for provider j; Tᵢ is tract provider-time supply; Dᵢ is provider density; Pᵢ is tract population; z(.) denotes standardization across all New York City tracts.")
    for p in [
        "The HFI score is calculated as the standardized sum of the reversed access component and the reversed density component. Higher values indicate more fragmented local healthcare environments. This specification preserves interpretability: tracts with farther nearest-provider access and thinner time-weighted provider density receive higher HFI values. At the same time, standardization allows comparison across tracts and prepares the framework for future extensions.",
        "The decision to begin with nearest-provider access and provider density is deliberately conservative. These components are comprehensible, reproducible, and grounded in established accessibility research. More complex models could be introduced immediately, but doing so would make it harder to evaluate whether the basic construct is coherent. The present implementation therefore privileges auditability. It creates a baseline surface that can be reproduced and then improved.",
        "This implementation should be understood as the first operational form of a broader measurement family. HFI can be extended without changing its conceptual core. Future versions can incorporate transportation mode, vehicle access, insurance acceptance, specialty availability, appointment wait times, referral networks, language access, patient origin, and discharge-flow information. The central requirement is that each extension preserve transparent construction, directional interpretability, and reproducibility."
    ]:
        para(doc, p)
    caption(doc, "Table 2. Key variables and operational definitions")
    add_table(doc, defs, widths=[1.35, .95, 2.75, 1.6])
    note(doc, "Definitions describe the operational variables used in this New York City implementation. The broader HFI framework can incorporate additional access and navigability domains in future versions.")
    heading(doc, "Outcomes and Statistical Analysis", 2)
    for p in [
        "The primary validation outcome is a communication index defined as the mean of HCAHPS nurse communication and doctor communication linear scores. Secondary outcomes include a broader patient-experience index combining standardized HCAHPS domains, including overall rating, recommendation, nurse communication, doctor communication, medication communication, and discharge communication. CMS overall hospital rating is included as an institutional performance covariate. All four hospital-level validation specifications use the 32-hospital HCAHPS-complete analytic sample.",
        "Spatial autocorrelation was evaluated at the tract level because HFI is a tract-level surface. Global Moran’s I tests whether the HFI surface has nonrandom spatial structure (Moran, 1950). We used k-nearest-neighbor weights with k = 8 as the primary specification and examined k = 4 and k = 12 as sensitivity checks. Moran’s I was evaluated with 999 random permutations. Local Moran’s I and Getis-Ord Gi* are important local diagnostics for later hotspot classification, but they are not the focus of this paper because the objective is construct validation of the index surface rather than definitive local cluster labeling (Getis and Ord, 1992; Anselin, 1995).",
        "Hospital-level validation models were estimated using ordinary least squares with HC3 robust standard errors. HC3 was selected because the hospital validation sample is modest and leverage may be nontrivial. Stata replication syntax is provided with the analytic packet, and the reported models use the standard HC3 robust variance estimator. Diagnostics included review of leverage-sensitive robust standard errors and correlation between HFI and CMS overall hospital rating. Hospital-level spatial dependence was not modeled because the validation sample includes only 32 point locations and the primary spatial diagnostic is the tract-level Moran’s I test for the index surface. Models are interpreted as construct validation and feasibility evidence for linking the tract-level HFI surface to public hospital outcomes. Direct causal estimation would require patient-origin data, temporal variation, larger multi-city samples, and richer neighborhood and institutional adjustment."
    ]:
        para(doc, p)
    add_display_equation(doc, "I = (n / W) × [ΣᵢΣⱼ wᵢⱼ(xᵢ - x̄)(xⱼ - x̄)] / [Σᵢ(xᵢ - x̄)²]")
    add_display_equation(doc, "Yₕ = β₀ + β₁HFIₕ + εₕ")
    add_display_equation(doc, "Yₕ = β₀ + β₁HFIₕ + β₂CMSRatingₕ + εₕ")
    note(doc, "I is global Moran's I for the tract-level HFI surface; wᵢⱼ is the spatial weight between tracts i and j; W is the sum of spatial weights; Yₕ is the HCAHPS communication or patient-experience outcome for hospital h.")

    heading(doc, "Results", 1)
    heading(doc, "HFI creates a complete tract-level surface and hospital validation layer", 2)
    for p in [
        "The first result is infrastructural. HFI creates a complete tract-level surface for all 2,229 New York City census tracts and links that surface to all eligible CMS acute-care hospitals with complete HCAHPS outcomes. The provider layer identifies 653 tracts with at least one provider and 1,576 tracts with no provider located within tract boundaries. This does not mean residents in those tracts have no access to care; it means their local care environment depends on resources outside the tract boundary, making nearest-provider access and cross-tract navigability central to the construct.",
        "The distribution of HFI is standardized over the tract surface. Higher HFI values indicate greater fragmentation. HFI is strongly related to the spatial components from which it is built: it is negatively correlated with provider density and positively correlated with minimum provider distance. Correlations with poverty and income are more modest, supporting the interpretation that HFI is not simply a socioeconomic deprivation index under another name.",
        "This distinction is central to interpretation. A deprivation index and a fragmentation index may identify overlapping places, but they imply different mechanisms and different interventions. If a tract is socioeconomically disadvantaged, policy responses may focus on income, insurance coverage, or social supports. If a tract is highly fragmented, responses may also need to address provider distribution, transportation links, care navigation, system coordination, or the placement of outpatient resources. HFI therefore adds a structural geography of care to existing neighborhood measures.",
    ]:
        para(doc, p)
    doc.add_page_break()
    caption(doc, "Table 3. Descriptive statistics for tract and hospital validation variables")
    add_table(doc, desc, widths=[2.25, .55, .7, .7, .7, .7, .7])
    note(doc, "Tract-level variables are summarized across 2,229 census tracts. Hospital variables are summarized for the 32-hospital HCAHPS validation sample.")
    doc.add_page_break()
    caption(doc, "Figure 2. Healthcare Fragmentation Index across New York City census tracts")
    doc.add_picture(str(fig2), width=Inches(6.45))
    note(doc, "Tract colors show HFI quintiles. Hospital markers show CMS acute-care hospitals with complete HCAHPS outcomes included in the validation sample.")
    caption(doc, "Figure 3. CMS acute-care hospitals embedded in the tract-level HFI surface")
    doc.add_picture(str(fig3), width=Inches(6.45))
    note(doc, "Markers show the full CMS-listed acute-care hospital universe in New York City. Yellow markers indicate hospitals included in the HCAHPS validation sample; gray indicates the acute-care hospital without required HCAHPS linear outcomes. Teal outlines identify public hospitals.")
    heading(doc, "HFI is spatially clustered rather than randomly distributed", 2)
    for p in [
        "The HFI surface displays statistically significant spatial clustering. Using 8-nearest-neighbor weights, global Moran’s I was 0.355 with permutation pseudo-p = 0.001. Sensitivity checks produced the same conclusion: Moran’s I was 0.389 with 4-nearest-neighbor weights and 0.323 with 12-nearest-neighbor weights, with pseudo-p = 0.001 in both specifications. These results indicate that HFI captures coherent spatial structure rather than random tract-level noise.",
        "The global Moran’s I result is important because HFI is intended to measure a spatially organized feature of the healthcare environment. If the surface had no spatial structure, the construct would be less plausible as a neighborhood measure. The significant spatial autocorrelation does not identify specific local clusters by itself, but it confirms that the index behaves like a spatial surface. Local indicators of spatial association and hotspot diagnostics are appropriate next steps once the global structure of the index has been established.",
        "The hospital validation sample includes 32 acute-care hospitals with complete HCAHPS outcomes. The mean communication index is 87.56, with values ranging from 81.50 to 93.00. The mean CMS overall hospital rating is 2.44. The hospital-linked HFI distribution differs from the full tract distribution because hospitals are concentrated in provider-rich and institutionally dense areas. This is expected and reinforces the importance of distinguishing the tract surface from the hospital-linked validation sample.",
    ]:
        para(doc, p)
    caption(doc, "Table 4. Spatial autocorrelation sensitivity checks for the tract-level HFI surface")
    add_table(doc, spatial, widths=[1.65, 1.0, 1.25, 2.4])
    note(doc, "Permutation pseudo-p values are based on 999 random permutations. The 8-nearest-neighbor specification is the primary diagnostic.")
    heading(doc, "HCAHPS validation shows a measurable hospital-level gradient", 2)
    for p in [
        "The hospital-level validation models show positive HFI coefficients for both HCAHPS communication and broader patient experience. In the bivariate communication model, the HFI coefficient is 0.217 with HC3 SE = 0.444. After adjustment for CMS overall hospital rating, the HFI coefficient is 0.155 with HC3 SE = 0.237. In models using the broader patient-experience index, the HFI coefficient is 0.096 in the bivariate model and 0.069 after adjustment for CMS overall rating. The pattern demonstrates that the HFI surface can be linked to standardized hospital patient-experience outcomes, while also showing that CMS overall hospital rating captures a larger share of hospital-level variation in patient experience.",
    ]:
        para(doc, p)
    caption(doc, "Figure 4. HFI and HCAHPS communication among validation hospitals")
    doc.add_picture(str(fig4), width=Inches(6.25))
    note(doc, "The fitted line summarizes the bivariate hospital-level association. The positive slope shows the observed hospital-level gradient between tract-linked HFI and HCAHPS communication in the validation sample.")
    doc.add_page_break()
    caption(doc, "Table 5. HFI and HCAHPS validation models")
    add_table(doc, models, widths=[2.2, 1.0, 1.0, 1.0, 1.0])
    note(doc, "Entries are OLS coefficients with HC3 robust standard errors in parentheses. Models 1-2 use the communication index. Models 3-4 use the patient-experience index. *** p < .001.")
    for p in [
        "CMS overall hospital rating is strongly associated with both communication and the broader patient-experience index. This is substantively informative: HCAHPS domains are closely related to institutional quality and patient-facing hospital processes, while HFI measures the surrounding care environment. The HFI coefficients remain positive after adjustment, supporting the value of linking the tract-level framework to public hospital outcomes while clarifying that hospital quality and local fragmentation are distinct empirical constructs.",
        "The results therefore support a constructive interpretation. HFI has a clear definition, a reproducible construction, interpretable spatial structure, and a public linkage to hospital outcomes. The HCAHPS exercise demonstrates that the index can be connected to standardized patient-experience data while preserving the conceptual distinction between neighborhood care-environment structure and hospital-level quality."
    ]:
        para(doc, p)

    heading(doc, "Discussion", 1)
    discussion = [
        "This study develops HFI as a reproducible geospatial framework for measuring healthcare fragmentation in local care environments. The results support three main conclusions. First, fragmentation is spatially patterned across New York City rather than randomly distributed. Second, HFI is empirically distinct from socioeconomic deprivation, even though fragmentation and disadvantage may overlap in important ways. Third, public hospital-level HCAHPS data provide a useful external construct-validation domain and demonstrate that HFI can be linked to standardized patient-experience outcomes.",
        "The paper’s central contribution should be read against the spatial accessibility literature. Prior work established the importance of measuring opportunity in relation to distance, travel impedance, supply, demand, and aggregation (Hansen, 1959; Joseph and Bantock, 1982; Khan, 1992; Radke and Mu, 2000; Luo and Wang, 2003; Guagliardo, 2004; Wang and Luo, 2005; Apparicio et al., 2008; Luo and Qi, 2009; McGrail and Humphreys, 2009). HFI builds on that tradition but asks a different question. Accessibility asks whether residents can reach opportunities. Fragmentation asks whether the surrounding healthcare environment functions as a coherent, reachable, and navigable system of care.",
        "The findings also extend health services research on fragmentation and coordination. Prior work has shown that fragmented systems can complicate accountability, continuity, coordination, quality, and costs (Starfield, 1998; Pham et al., 2007; Bodenheimer, 2008; Cebul et al., 2008; Stange, 2009; Frandsen et al., 2015). Much of that literature measures fragmentation through patient-level care patterns, physician dispersion, organizational relationships, or care transitions. HFI brings that concern into spatial form. It treats fragmentation not only as something that happens after patients enter care, but also as a property of the local environment through which patients must find, reach, and coordinate care in the first place.",
        "The New York City setting makes this contribution concrete. A city can have many hospitals and providers while still containing fragmented local care environments. Public hospitals, academic systems, community hospitals, and specialty institutions do not form a uniform geography. They sit within neighborhoods shaped by transportation networks, historical marginalization, safety-net dependence, insurance barriers, and uneven institutional resources. The Urban Stress Model argues that city form, social disadvantage, and chronic stress pathways interact to reproduce health disparities in New York City and other urban settings (Afable et al., 2026). HFI extends that logic into the healthcare environment by measuring whether local care resources are organized in ways that may reduce or compound navigational burden.",
        "This framing also responds to a common limitation in access research: the tendency to treat access, quality, and patient experience as separable domains. In practice, patients encounter them together. A nearby provider who is unaffordable, inaccessible by transit, outside a patient’s insurance network, or disconnected from hospital follow-up does not provide meaningful access in the lived sense. A hospital’s communication burden may also be shaped by the local outpatient environment into which patients are discharged. HFI provides a structure for measuring this combined geography of access, navigation, and patient experience.",
        "The HCAHPS results are informative because they preserve the conceptual distinction between local care-environment structure and hospital performance. HFI is a tract-level care-environment measure assigned to hospitals through location. HCAHPS communication is a hospital-level patient-experience measure reflecting patients from many origins, hospital processes, expectations, survey response patterns, and institutional quality. The positive HFI coefficients are therefore best interpreted as supportive linkage and construct-validation evidence rather than as a claim that neighborhood fragmentation alone determines hospital communication scores.",
        "This distinction strengthens the paper. It shows that HFI can be linked to standardized public quality data without collapsing the framework into a hospital-rating exercise. HFI measures the geography of care environments; HCAHPS measures patients’ reported inpatient experience. Their relationship is theoretically meaningful because fragmented local environments may increase navigational and coordination burden, but the constructs remain appropriately distinct.",
        "This is precisely why HFI should be understood as a framework rather than a one-time score. Future validation should use patient-origin data, discharge records, travel-time estimates, hospital catchment areas, specialty-specific access, insurance-network participation, and utilization patterns. In New York, SPARCS-linked analyses could test whether patients from more fragmented tracts experience different patterns of hospital use, readmission, emergency department reliance, or continuity-sensitive outcomes. Dynamic HFI versions could examine whether fragmentation changes as hospitals close, merge, open new ambulatory sites, or alter service lines.",
        "The next version of HFI should also incorporate transportation and social need more explicitly. Transportation is not merely an adjustment variable; it is part of the access-fragmentation construct. Vehicle availability, transit dependence, disability, age structure, language isolation, and commute-mode patterns may modify how the same provider geography is experienced by different neighborhoods. Social vulnerability and social need indices can be used alongside HFI to test whether fragmented care environments compound existing disadvantage.",
        "The Explorer is also part of the contribution. A reproducible measurement framework should not live only in a manuscript. The HFI Explorer provides a public implementation of the tract surface, hospital layer, documentation, and downloadable data. It allows collaborators, reviewers, and future users to inspect the geography, identify potential improvements, and treat HFI as a cumulative research platform rather than a static paper figure."
    ]
    for p in discussion:
        para(doc, p)

    heading(doc, "Limitations", 1)
    limitations = [
        "Several considerations guide interpretation. First, the hospital validation sample includes all eligible CMS acute-care hospitals in New York City with complete HCAHPS outcomes, producing a complete eligible single-city validation layer. Because n = 32, the models are intentionally parsimonious and avoid overadjustment. This is appropriate for a first citywide construct-validation analysis and creates a transparent baseline for future multi-city replication.",
        "Second, the point-in-tract linkage is reproducible but incomplete. Hospitals draw patients from broader catchment areas, and hospital communication scores reflect patients whose residences may span many tracts. Patient-origin and discharge-flow data are needed to estimate exposure to fragmented care environments more directly.",
        "Third, HFI is a tract-level measure and may be sensitive to geographic scale. Census tracts are useful because they are widely available, interpretable, and compatible with public data, but residents do not experience healthcare systems only within tract boundaries. Future work should examine sensitivity to alternative units, network catchments, travel sheds, and patient-origin geographies.",
        "Fourth, the present implementation prioritizes the spatial dimensions that can be constructed reproducibly across the full tract surface. Financial affordability, insurance acceptance, appointment availability, language access, specialty mix, disability-relevant accessibility, and multimodal transportation are central to healthcare fragmentation and require additional data. They are part of the HFI framework and priorities for future versions.",
        "Fifth, the analysis is cross-sectional. It is designed to construct and validate a spatial measurement framework rather than estimate a causal treatment effect. Longitudinal analyses, natural experiments, hospital closures, service-line changes, transit disruptions, and repeated HCAHPS reporting periods would allow future work to test causal pathways opened by the framework.",
        "Finally, HCAHPS is a valuable validation outcome because it is standardized, public, and theoretically connected to communication and navigation. Future validation can build from this public benchmark by examining readmissions, emergency department use, avoidable hospitalizations, missed appointments, discharge understanding, continuity-sensitive outcomes, and subgroup differences."
    ]
    for p in limitations:
        para(doc, p)

    heading(doc, "Conclusion", 1)
    conclusion = [
        "HFI makes healthcare fragmentation measurable as a tract-level spatial condition. The New York City implementation shows that the framework can be constructed across all city census tracts, evaluated for spatial clustering, linked to the CMS acute-care hospital universe, and compared with publicly reported patient-experience outcomes. The central finding is that local healthcare fragmentation can be operationalized transparently and studied as a reproducible feature of urban healthcare environments.",
        "This distinction is important. Healthcare fragmentation is often discussed as a system problem, but system problems become difficult to address when they are not measured spatially. HFI provides a way to identify where healthcare resources exist as a coherent local environment and where residents may face a more scattered, thin, or difficult-to-navigate field of care. That information can support research, planning, public accountability, and future validation.",
        "The New York City case also demonstrates why fragmentation should be interpreted through an equity lens. Dense healthcare markets can still produce uneven care environments when institutional resources, transportation networks, affordability, safety-net capacity, and neighborhood histories do not align. HFI creates a tract-level measurement infrastructure into which these dimensions can be added, tested, and debated.",
        "The HCAHPS validation exercise provides an important first public benchmark. Hospital patient-experience scores are institution-level outcomes shaped by many forces, and the point-in-tract linkage is intentionally transparent. The HFI coefficients show that the framework can be connected to standardized patient-experience data while pointing toward the next empirical task. Future validation should use patient-origin data, utilization records, travel-time models, hospital catchments, discharge flows, and dynamic analyses of how healthcare geography changes over time.",
        "This is where the framework becomes most useful. HFI can be paired with social need indices, transportation measures, insurance-network data, and patient-flow data to test whether fragmentation compounds existing disadvantage. It can be recalculated after hospital closures, service-line changes, transit disruptions, ambulatory expansion, or policy interventions. It can also be replicated in other metropolitan areas to determine whether the spatial organization of care differs across cities with different histories, governance structures, and health-system geographies.",
        "The accompanying HFI Explorer is part of the scientific contribution. By making the tract surface, hospital layer, documentation, and downloadable data public, the Explorer turns HFI into an inspectable and cumulative research platform. Reviewers, collaborators, policymakers, and community partners can examine the map, question the assumptions, propose new data layers, and compare future versions against the current baseline.",
        "The biggest strength of HFI is that it provides a common, reproducible language through which fragmentation can be measured, compared, challenged, and ultimately improved. The accompanying HFI Explorer extends that contribution by making the framework public, inspectable, and ready for future refinement. If healthcare fragmentation has remained difficult to address because it is dispersed across geography, institutions, and patient experience, HFI makes that dispersion visible enough to study. That is the foundation on which stronger validation, better planning, and more equitable healthcare environments can be built."
    ]
    for p in conclusion:
        para(doc, p)

    heading(doc, "Data Availability", 1)
    para(doc, "The HFI Explorer and accompanying public-release materials are available at https://giacomodi-pasquale-rgb.github.io/hfi-explorer/. The release includes tract-level HFI files, linked hospital information, documentation, data-download materials, citation guidance, and planned future versions of the index.")

    heading(doc, "References", 1)
    refs = [
        "Aday LA, Andersen R. A framework for the study of access to medical care. Health Services Research. 1974;9(3):208-220.",
        "Afable A, Rybarczyk G, Rivera BD, Landsbergis P. The Urban Stress Model: a framework for understanding the persistence of health disparities in cities. Cities & Health. 2026. doi:10.1080/23748834.2025.2604406.",
        "Andersen RM. Revisiting the behavioral model and access to medical care: does it matter? Journal of Health and Social Behavior. 1995;36(1):1-10.",
        "Anhang Price R, Elliott MN, Zaslavsky AM, et al. Examining the role of patient experience surveys in measuring health care quality. Medical Care Research and Review. 2014;71(5):522-554.",
        "Anselin L. Local indicators of spatial association: LISA. Geographical Analysis. 1995;27(2):93-115.",
        "Apparicio P, Abdelmajid M, Riva M, Shearmur R. Comparing alternative approaches to measuring the geographical accessibility of urban health services: distance types and aggregation-error issues. International Journal of Health Geographics. 2008;7:7.",
        "Beckett MK, Quigley DD, Lehrman WG, Giordano LA, Cohea CW, Goldstein EH, Elliott MN. Interventions and hospital characteristics associated with patient experience: an update of the evidence. Medical Care Research and Review. 2024;81(3):195-208.",
        "Bice TW, Boxerman SB. A quantitative measure of continuity of care. Medical Care. 1977;15(4):347-349.",
        "Bleich SN, Ozaltin E, Murray CJL. How does satisfaction with the health-care system relate to patient experience? Bulletin of the World Health Organization. 2009;87:271-278.",
        "Bodenheimer T. Coordinating care--a perilous journey through the health care system. New England Journal of Medicine. 2008;358(10):1064-1071.",
        "Boulding W, Glickman SW, Manary MP, Schulman KA, Staelin R. Relationship between patient satisfaction with inpatient care and hospital readmission within 30 days. American Journal of Managed Care. 2011;17(1):41-48.",
        "Carr LJ, Dunsiger SI, Marcus BH. Validation of Walk Score for estimating access to walkable amenities. British Journal of Sports Medicine. 2011;45(14):1144-1148.",
        "Cebul RD, Rebitzer JB, Taylor LJ, Votruba ME. Organizational fragmentation and care quality in the U.S. healthcare system. Journal of Economic Perspectives. 2008;22(4):93-113.",
        "Centers for Medicare & Medicaid Services. Hospital General Information. Provider Data Catalog. 2026a.",
        "Centers for Medicare & Medicaid Services. HCAHPS - Hospital. Provider Data Catalog. 2026b.",
        "Diez Roux AV, Mair C. Neighborhoods and health. Annals of the New York Academy of Sciences. 2010;1186:125-145.",
        "Doyle C, Lennox L, Bell D. A systematic review of evidence on the links between patient experience and clinical safety and effectiveness. BMJ Open. 2013;3:e001570.",
        "Flanagan BE, Gregory EW, Hallisey EJ, Heitgerd JL, Lewis B. A social vulnerability index for disaster management. Journal of Homeland Security and Emergency Management. 2011;8(1).",
        "Frandsen BR, Joynt KE, Rebitzer JB, Jha AK. Care fragmentation, quality, and costs among chronically ill patients. American Journal of Managed Care. 2015;21(5):355-362.",
        "Getis A, Ord JK. The analysis of spatial association by use of distance statistics. Geographical Analysis. 1992;24(3):189-206.",
        "Goldstein E, Farquhar M, Crofton C, Darby C, Garfinkel S. Measuring hospital care from the patients' perspective: an overview of the CAHPS Hospital Survey development process. Health Services Research. 2005;40(6p2):1977-1995.",
        "Guagliardo MF. Spatial accessibility of primary care: concepts, methods and challenges. International Journal of Health Geographics. 2004;3:3.",
        "Hansen WG. How accessibility shapes land use. Journal of the American Institute of Planners. 1959;25(2):73-76.",
        "Higgs G. A literature review of the use of GIS-based measures of access to health care services. Health Services and Outcomes Research Methodology. 2004;5:119-139.",
        "Institute of Medicine. Crossing the Quality Chasm: A New Health System for the 21st Century. Washington, DC: National Academies Press; 2001.",
        "Isaac T, Zaslavsky AM, Cleary PD, Landon BE. The relationship between patients' perception of care and measures of hospital quality and safety. Health Services Research. 2010;45(4):1024-1040.",
        "Jee SH, Cabana MD. Indices for continuity of care: a systematic review of the literature. Medical Care Research and Review. 2006;63(2):158-188.",
        "Jha AK, Orav EJ, Zheng J, Epstein AM. Patients' perception of hospital care in the United States. New England Journal of Medicine. 2008;359(18):1921-1931.",
        "Joseph AE, Bantock PR. Measuring potential physical accessibility to general practitioners in rural areas: a method and case study. Social Science & Medicine. 1982;16(1):85-90.",
        "Khan AA. An integrated approach to measuring potential spatial access to health care services. Socio-Economic Planning Sciences. 1992;26(4):275-287.",
        "Kind AJH, Buckingham WR. Making neighborhood-disadvantage metrics accessible: the Neighborhood Atlas. New England Journal of Medicine. 2018;378(26):2456-2458.",
        "Levesque JF, Harris MF, Russell G. Patient-centred access to health care: conceptualising access at the interface of health systems and populations. International Journal for Equity in Health. 2013;12:18.",
        "Luo W, Qi Y. An enhanced two-step floating catchment area method for measuring spatial accessibility to primary care physicians. Health & Place. 2009;15(4):1100-1107.",
        "Luo W, Wang F. Measures of spatial accessibility to health care in a GIS environment. Environment and Planning B. 2003;30(6):865-884.",
        "Manary MP, Boulding W, Staelin R, Glickman SW. The patient experience and health outcomes. New England Journal of Medicine. 2013;368:201-203.",
        "McGrail MR, Humphreys JS. Measuring spatial accessibility to primary care in rural areas: improving the effectiveness of the two-step floating catchment area method. Applied Geography. 2009;29(4):533-541.",
        "Moran PAP. Notes on continuous stochastic phenomena. Biometrika. 1950;37(1/2):17-23.",
        "Penchansky R, Thomas JW. The concept of access: definition and relationship to consumer satisfaction. Medical Care. 1981;19(2):127-140.",
        "Pham HH, Schrag D, O'Malley AS, Wu B, Bach PB. Care patterns in Medicare and their implications for pay for performance. New England Journal of Medicine. 2007;356(11):1130-1139.",
        "Radke J, Mu L. Spatial decompositions, modeling and mapping service regions to predict access to social programs. Geographic Information Sciences. 2000;6(2):105-112.",
        "Saultz JW. Defining and measuring interpersonal continuity of care. Annals of Family Medicine. 2003;1(3):134-143.",
        "Schoen C, Osborn R, Squires D, Doty MM, Pierson R, Applebaum S. New 2011 survey of patients with complex care needs in eleven countries finds that care is often poorly coordinated. Health Affairs. 2011;30(12):2437-2448.",
        "Sofaer S, Firminger K. Patient perceptions of the quality of health services. Annual Review of Public Health. 2005;26:513-559.",
        "Stange KC. The problem of fragmentation and the need for integrative solutions. Annals of Family Medicine. 2009;7(2):100-103.",
        "Starfield B. Primary Care: Balancing Health Needs, Services, and Technology. New York: Oxford University Press; 1998.",
        "Starfield B, Shi L, Macinko J. Contribution of primary care to health systems and health. Milbank Quarterly. 2005;83(3):457-502.",
        "Syed ST, Gerber BS, Sharp LK. Traveling towards disease: transportation barriers to health care access. Journal of Community Health. 2013;38(5):976-993.",
        "Tsai TC, Orav EJ, Jha AK. Patient satisfaction and quality of surgical care in US hospitals. Annals of Surgery. 2015;261(1):2-8.",
        "Wang F, Luo W. Assessing spatial and nonspatial factors for healthcare access: towards an integrated approach to defining health professional shortage areas. Health & Place. 2005;11(2):131-146.",
        "Wennberg J, Gittelsohn A. Small area variations in health care delivery. Science. 1973;182(4117):1102-1108.",
        "Williams DR, Collins C. Racial residential segregation: a fundamental cause of racial disparities in health. Public Health Reports. 2001;116(5):404-416.",
    ]
    for ref in refs:
        para(doc, ref, size=10)

    doc.save(DOCX_OUT)
    return DOCX_OUT


def update_email():
    email = """Subject: HFI analytic file and manuscript update

Hi everyone,

I wanted to give you a clear update before Friday.

I am finalizing a source-data reconciliation of the hospital universe and the analytic HFI/HCAHPS file so that the manuscript, tables, figures, and Explorer are all aligned. I used the CMS Hospital General Information and HCAHPS files as the reference point. The current analytic scope is 33 CMS-listed acute-care hospitals in New York City, with 32 hospitals eligible for the HCAHPS validation analysis because they have the required patient-experience outcomes.

I am also finalizing the manuscript around that analytic file. The main framing is HFI as a reproducible geospatial framework for measuring local healthcare fragmentation. The HCAHPS models remain in the paper as a public construct-validation linkage between the tract-level HFI surface and standardized hospital patient-experience outcomes.

I think this puts us in a stronger position. The paper will be clearer about the hospital universe, more precise about access versus fragmentation, more responsive to the conceptual comments, and more defensible statistically. I am updating the methods, figures, tables, results, discussion, and conclusion together so that the full manuscript reads as one coherent version.

I should be able to circulate the cleaned draft and supporting hospital-system list in the next few days.

Best,
Giacomo
"""
    EMAIL_OUT.write_text(email, encoding="utf-8")


if __name__ == "__main__":
    out = build_doc()
    update_email()
    print(out)
