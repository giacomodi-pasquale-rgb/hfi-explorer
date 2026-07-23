from pathlib import Path
from math import erf, sqrt

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/giacomo/Documents/Codex/2026-07-07/is")
OUT = ROOT / "outputs/hfi_surface_rebuild_2026_07_22/word_ready_tables_corrected_hfi"
OUT.mkdir(parents=True, exist_ok=True)

ANALYSIS = ROOT / "outputs/hfi_surface_rebuild_2026_07_22/nyc_hcahps_validation_analysis_file_corrected_hfi_v1_1.csv"
UNIVERSE = ROOT / "outputs/hospital_rebuild_2026_07_22/derived/nyc_acute_care_hospital_universe_33.csv"


def normal_p(t):
    return 2 * (1 - 0.5 * (1 + erf(abs(t) / sqrt(2))))


def fmt(x, digits=2):
    if pd.isna(x):
        return ""
    return f"{x:.{digits}f}"


def star(p):
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    if p < 0.10:
        return "+"
    return ""


def numeric(df, cols):
    for col in cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


def ols_hc3(df, y, xs):
    d = df[[y] + xs].dropna().copy()
    yv = d[y].to_numpy(float)
    X = np.column_stack([np.ones(len(d))] + [d[x].to_numpy(float) for x in xs])
    names = ["Intercept"] + xs
    inv = np.linalg.inv(X.T @ X)
    beta = inv @ (X.T @ yv)
    fitted = X @ beta
    resid = yv - fitted
    hat = np.sum(X * (X @ inv), axis=1)
    omega = (resid / (1 - hat)) ** 2
    vcov = inv @ (X.T @ (omega[:, None] * X)) @ inv
    se = np.sqrt(np.diag(vcov))
    sst = float(((yv - yv.mean()) ** 2).sum())
    ssr = float((resid ** 2).sum())
    r2 = 1 - ssr / sst if sst else np.nan
    rows = {}
    for name, b, s in zip(names, beta, se):
        t = b / s if s else np.nan
        rows[name] = {"coef": b, "se": s, "t": t, "p": normal_p(t)}
    return {"n": len(d), "r2": r2, "terms": rows}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value, bold=False, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(value))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_rows=1):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)
            if row_idx < header_rows:
                set_cell_shading(cell, "EAF3F5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_caption(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)


def add_note(doc, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Note. " + note)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(10)


def add_df_table(doc, df, widths=None):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], val, size=10, align=align)
    style_table(table)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def build_tables():
    hosp = pd.read_csv(ANALYSIS, dtype=str)
    universe = pd.read_csv(UNIVERSE, dtype=str)
    num_cols = [
        "fragmentation_index", "communication_index", "patient_experience_index",
        "hospital_overall_rating", "number_of_completed_surveys",
        "survey_response_rate_percent", "hcahps_nurse_comm_linear",
        "hcahps_doctor_comm_linear", "hcahps_medicine_comm_linear",
        "hcahps_discharge_linear", "hcahps_overall_linear", "hcahps_recommend_linear",
    ]
    hosp = numeric(hosp, num_cols)

    county_order = ["BRONX", "KINGS", "NEW YORK", "QUEENS", "RICHMOND"]
    t1 = hosp.groupby("county").agg(
        hospitals=("facility_id", "count"),
        public_hospitals=("public_private_designation", lambda s: int((s == "Public").sum())),
        private_nonprofit_hospitals=("public_private_designation", lambda s: int((s != "Public").sum())),
        mean_hfi=("fragmentation_index", "mean"),
        mean_communication=("communication_index", "mean"),
    ).reindex(county_order).reset_index()
    t1["mean_hfi"] = t1["mean_hfi"].map(lambda x: fmt(x, 2))
    t1["mean_communication"] = t1["mean_communication"].map(lambda x: fmt(x, 1))
    t1.columns = ["Borough/county", "Hospitals", "Public", "Private nonprofit", "Mean HFI", "Mean communication"]

    desc_vars = [
        ("Healthcare Fragmentation Index", "fragmentation_index"),
        ("Communication index", "communication_index"),
        ("Patient-experience index", "patient_experience_index"),
        ("CMS overall hospital rating", "hospital_overall_rating"),
        ("Completed HCAHPS surveys", "number_of_completed_surveys"),
        ("Survey response rate (%)", "survey_response_rate_percent"),
    ]
    t2_rows = []
    for label, col in desc_vars:
        s = hosp[col].dropna()
        t2_rows.append({
            "Variable": label,
            "N": int(s.count()),
            "Mean": fmt(s.mean(), 2),
            "SD": fmt(s.std(ddof=1), 2),
            "Min": fmt(s.min(), 2),
            "Median": fmt(s.median(), 2),
            "Max": fmt(s.max(), 2),
        })
    t2 = pd.DataFrame(t2_rows)

    specs = [
        ("Model 1", "communication_index", ["fragmentation_index"]),
        ("Model 2", "communication_index", ["fragmentation_index", "hospital_overall_rating"]),
        ("Model 3", "patient_experience_index", ["fragmentation_index"]),
        ("Model 4", "patient_experience_index", ["fragmentation_index", "hospital_overall_rating"]),
    ]
    models = [(name, ols_hc3(hosp, y, xs)) for name, y, xs in specs]
    terms = [
        ("Healthcare Fragmentation Index", "fragmentation_index"),
        ("CMS overall hospital rating", "hospital_overall_rating"),
        ("Constant", "Intercept"),
    ]
    t3_rows = []
    for label, key in terms:
        row = {"Predictor": label}
        se_row = {"Predictor": ""}
        for name, model in models:
            if key in model["terms"]:
                item = model["terms"][key]
                row[name] = fmt(item["coef"], 3) + star(item["p"])
                se_row[name] = "(" + fmt(item["se"], 3) + ")"
            else:
                row[name] = ""
                se_row[name] = ""
        t3_rows.append(row)
        t3_rows.append(se_row)
    n_row = {"Predictor": "N"}
    r2_row = {"Predictor": "R-squared"}
    for name, model in models:
        n_row[name] = str(model["n"])
        r2_row[name] = fmt(model["r2"], 3)
    t3_rows.extend([n_row, r2_row])
    t3 = pd.DataFrame(t3_rows)

    tA = universe[["facility_id", "facility_name", "county", "hospital_system", "included_in_hcahps_validation", "hcahps_exclusion_reason"]].copy()
    tA["included_in_hcahps_validation"] = tA["included_in_hcahps_validation"].map(lambda x: "Yes" if str(x).lower() == "true" else "No")
    tA["hcahps_exclusion_reason"] = tA["hcahps_exclusion_reason"].fillna("")
    tA.columns = ["CMS ID", "Hospital", "County", "System", "Included", "Reason if excluded"]
    return t1, t2, t3, tA


def make_docx(t1, t2, t3, tA):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Word-Ready Tables for Rebuilt 32-Hospital HFI/HCAHPS Analysis")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CMS-listed NYC acute-care hospitals with complete HCAHPS outcomes; n = 32.")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)

    add_caption(doc, "Table 1. HCAHPS validation hospitals by borough/county")
    add_df_table(doc, t1)
    add_note(doc, "The analytic sample includes CMS-listed acute-care hospitals in New York City with complete HCAHPS patient-experience outcomes and successful linkage to the tract-level HFI surface.")

    add_caption(doc, "Table 2. Descriptive statistics for rebuilt hospital validation sample")
    add_df_table(doc, t2, widths=[2.3, 0.55, 0.8, 0.8, 0.8, 0.8, 0.8])
    add_note(doc, "HFI values are tract-level scores assigned to each hospital by spatial linkage. Patient-experience values are from CMS HCAHPS hospital files.")

    doc.add_page_break()
    add_caption(doc, "Table 3. HFI and HCAHPS validation models")
    add_df_table(doc, t3, widths=[2.3, 1.2, 1.2, 1.2, 1.2])
    add_note(doc, "Entries are OLS coefficients with HC3 robust standard errors in parentheses. + p < .10, * p < .05, ** p < .01, *** p < .001 using normal approximation. Models 1-2 use the communication index; Models 3-4 use the patient-experience index.")

    doc.add_page_break()
    add_caption(doc, "Appendix Table A1. CMS NYC acute-care hospital universe and HCAHPS inclusion")
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    add_df_table(doc, tA, widths=[0.7, 2.5, 0.8, 2.1, 0.7, 2.2])
    add_note(doc, "NY Eye and Ear Infirmary of Mount Sinai is retained in the CMS acute-care universe but excluded from HCAHPS validation because required HCAHPS linear outcomes are unavailable.")

    path = OUT / "HFI_32_hospital_corrected_hfi_word_ready_tables.docx"
    doc.save(path)
    return path


def make_stata_text(t1, t2, t3, tA):
    lines = []
    lines.append("TABLE 1. HCAHPS validation hospitals by borough/county")
    lines.append(t1.to_string(index=False))
    lines.append("")
    lines.append("TABLE 2. Descriptive statistics for rebuilt hospital validation sample")
    lines.append(t2.to_string(index=False))
    lines.append("")
    lines.append("TABLE 3. HFI and HCAHPS validation models")
    lines.append(t3.to_string(index=False))
    lines.append("")
    lines.append("APPENDIX TABLE A1. CMS NYC acute-care hospital universe and HCAHPS inclusion")
    lines.append(tA.to_string(index=False))
    path = OUT / "HFI_32_hospital_corrected_hfi_stata_style_tables.txt"
    path.write_text("\n".join(lines))
    return path


def main():
    t1, t2, t3, tA = build_tables()
    for name, table in [("table1_word_ready.csv", t1), ("table2_word_ready.csv", t2), ("table3_word_ready.csv", t3), ("appendix_tableA1_word_ready.csv", tA)]:
        table.to_csv(OUT / name, index=False)
    docx = make_docx(t1, t2, t3, tA)
    txt = make_stata_text(t1, t2, t3, tA)
    print(docx)
    print(txt)


if __name__ == "__main__":
    main()
